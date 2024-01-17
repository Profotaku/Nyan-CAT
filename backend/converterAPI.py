import base64
import os
import re
from io import BytesIO

import docx
from bs4 import BeautifulSoup
from docx.shared import Inches
from docx.oxml import OxmlElement
from htmldocx import HtmlToDocx
from docx.oxml.ns import qn

class ConverterAPI:
    def __init__(self):
        self.new_parser = HtmlToDocx()

    def insert_image_from_base64(self, doc, base64_str, ParaPosition, imgisbefore):
        image_data = base64.b64decode(base64_str)
        image_stream = BytesIO(image_data)
        paragraph = doc.paragraphs[ParaPosition]

        def find_image_run(paragraph):
            for run in paragraph.runs:
                if run.element.xpath('.//w:drawing'):
                    return run
            return None

        image_run = find_image_run(paragraph)

        if image_run:
            # Supprimer l'ancienne image
            for element in image_run.element.getchildren():
                if element.tag.endswith('drawing'):
                    image_run.element.remove(element)

        if imgisbefore:
            if not image_run or image_run != paragraph.runs[0]:
                # Créer un nouveau run au début si l'ancien run avec image n'est pas le premier
                new_run = OxmlElement('w:r')
                paragraph._p.insert(0, new_run)
                new_run = paragraph.runs[0]
                new_run.add_picture(image_stream, width=Inches(6))
            else:
                # L'ancien run avec image est le premier, ajouter l'image ici
                image_run.add_picture(image_stream, width=Inches(6))
        else:
            if not image_run or image_run != paragraph.runs[-1]:
                # Ajouter un nouveau run à la fin si l'ancien run avec image n'est pas le dernier
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6))
            else:
                # L'ancien run avec image est le dernier, ajouter l'image ici
                image_run.add_picture(image_stream, width=Inches(6))

    def ParaHtmlToDocx(self, Html: str, Docx: docx.document, ParaPosition: int, SaveName: str):
        soup = BeautifulSoup(Html, 'html.parser')
        p_tag = soup.find('p')
        img_tag = soup.find('img')
        if img_tag:
            #remove all img tags
            Html = re.sub(r'<img.*?>', '', Html)
        tempdocx = self.new_parser.parse_html_string(Html)  # docx file with only the html conversion
        tempdocx.save("temp.docx")
        #determine if img is before or after the text :
        imgisbefore = True
        imgisbefore = p_tag.contents and p_tag.contents[0] == img_tag
        for ParaInHtml in range(0, len(tempdocx.paragraphs)):
            if len(tempdocx.paragraphs) + ParaPosition > len(Docx.paragraphs):
                Docx.add_paragraph()
            Docx.paragraphs[ParaPosition + ParaInHtml].text = ""
            Docx.paragraphs[ParaPosition + ParaInHtml].alignment = tempdocx.paragraphs[ParaInHtml].alignment
            Docx.paragraphs[ParaPosition + ParaInHtml].style = tempdocx.paragraphs[ParaInHtml].style
            Docx.paragraphs[ParaPosition + ParaInHtml].style.name = tempdocx.paragraphs[ParaInHtml].style.name
            Docx.paragraphs[ParaPosition + ParaInHtml].style.base_style = tempdocx.paragraphs[ParaInHtml].style.base_style
            Docx.paragraphs[ParaPosition + ParaInHtml].style.priority = tempdocx.paragraphs[ParaInHtml].style.priority
            Docx.paragraphs[ParaPosition + ParaInHtml].style.style_id = tempdocx.paragraphs[ParaInHtml].style.style_id
            Docx.paragraphs[ParaPosition + ParaInHtml].paragraph_format.left_indent = tempdocx.paragraphs[ParaInHtml].paragraph_format.left_indent
            Docx.paragraphs[ParaPosition + ParaInHtml].paragraph_format.right_indent = tempdocx.paragraphs[ParaInHtml].paragraph_format.right_indent

            for idx, run in enumerate(tempdocx.paragraphs[ParaInHtml].runs):
                if len(tempdocx.paragraphs[ParaInHtml].runs) > len(Docx.paragraphs[ParaPosition + ParaInHtml].runs):
                    Docx.paragraphs[ParaPosition + ParaInHtml].add_run()
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].text = run.text
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].bold = run.bold
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].italic = run.italic
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].underline = run.underline
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].font.color.rgb = run.font.color.rgb
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].style.name = run.style.name
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].font.name = run.font.name
                Docx.paragraphs[ParaPosition + ParaInHtml].runs[idx].font.size = run.font.size

            images = soup.find_all('img')
            if len(images) > 0:
                img = images[0]
                img_data = img['src'].split('base64,')[-1]
                self.insert_image_from_base64(Docx, img_data, ParaPosition, imgisbefore)
            Docx.save(SaveName)


    def extract_and_get_img_html(self, run, doc_part):
        # Extraction de l'image
        image_html = ''
        inline_shapes = run.element.xpath('.//a:blip/../..')
        for inline_shape in inline_shapes:
            blip = inline_shape.xpath('.//a:blip')[0]
            r_id = blip.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed']
            image_part = doc_part.related_parts[r_id]
            image_ext = os.path.splitext(image_part.partname)[1]
            # Encodage de l'image en base64 pour intégration directe dans le HTML
            image_base64 = base64.b64encode(image_part._blob).decode('utf-8')
            image_html += '<img src="data:image/{};base64,{}" style="width:auto;height:auto;"/>'.format(image_ext.strip('.'), image_base64)
        return image_html
    def ParaDocxToHtml(self, Docx: docx.document, ParaPosition: int):
        if ParaPosition >= len(Docx.paragraphs):
            return ""
        if ParaPosition < 0:
            return ""
        ParatoConvert = Docx.paragraphs[ParaPosition]
        style = ParatoConvert.style.name
        valclass = ""
        if style == "Normal":
            tag = 'p'
        else:
            tag = re.sub(r'[a-z- ]+', '', style).lower()
        if tag == 't':
            tag = 'h1'
        elif "p" in tag:
            tag = "p"
        if ParatoConvert.alignment or ParatoConvert.paragraph_format.left_indent or ParatoConvert.style.paragraph_format.alignment or ParatoConvert.style.paragraph_format.left_indent:
            valclass = valclass + " class=\""
            if str(ParatoConvert.alignment) == "CENTER (1)" or str(ParatoConvert.style.paragraph_format.alignment) == "CENTER (1)":
                valclass = valclass + "ql-align-center"
            elif str(ParatoConvert.alignment) == "JUSTIFY (3)" or str(ParatoConvert.style.paragraph_format.alignment) == "JUSTIFY (3)":
                valclass = valclass + "ql-align-justify"
            elif str(ParatoConvert.alignment) == "RIGHT (2)" or str(ParatoConvert.style.paragraph_format.alignment) == "RIGHT (2)":
                valclass = valclass + "ql-align-right"
        if (ParatoConvert.paragraph_format.left_indent or ParatoConvert.style.paragraph_format.left_indent) and (ParatoConvert.alignment or ParatoConvert.style.paragraph_format.alignment):
            valclass = valclass + " "
        if ParatoConvert.paragraph_format.left_indent:
            num = round(float(ParatoConvert.paragraph_format.left_indent.cm)*2.37/3*4/5.5)
            valclass = valclass + "ql-indent-" + str(num)
        if ParatoConvert.paragraph_format.left_indent or ParatoConvert.alignment or ParatoConvert.style.paragraph_format.alignment or ParatoConvert.style.paragraph_format.left_indent:
            valclass = valclass + "\""
        Html = '<' + tag + valclass + '>'
        for run in ParatoConvert.runs:
            if run.text != '':
                if run.text[0] == '\n':
                    Html = Html + "</" + tag + "><" + tag + ">"
            if run.font.color.rgb is not None:
                RGB = tuple(int(str(run.font.color.rgb)[i:i + 2], 16) for i in (0, 2, 4))
                Html = Html + "<span style=\"color: rgb(" + str(RGB[0]) + ", " + str(RGB[1]) + ", " + str(RGB[2]) + ");\">"
            if run.italic:
                Html = Html + "<em>"
            if run.bold:
                Html = Html + "<strong>"
            if run.underline:
                Html = Html + "<u>"
            Html = Html + run.text
            if run.underline:
                Html = Html + "/<u>"
            if run.bold:
                Html = Html + "</strong>"
            if run.italic:
                Html = Html + "</em>"
            if run.font.color.rgb is not None:
                Html = Html + "</span>"
            if run.element.xpath('.//a:blip'):
                Html = Html + self.extract_and_get_img_html(run, Docx.part)
        Html = Html + '</' + tag + '>'
        return Html
